import os
import sqlite3
import whisper

WHISPER_MODEL = "base"  # Change model if needed
DB_FILE = "transcripts.db"

# Load Whisper model
model = whisper.load_model(WHISPER_MODEL)

# Ensure DB exists
conn = sqlite3.connect(DB_FILE)
c = conn.cursor()
c.execute("""
CREATE TABLE IF NOT EXISTS transcripts (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    filename TEXT UNIQUE,
    transcript TEXT
)
""")
conn.commit()
conn.close()

def transcribe_and_store(audio_path):
    """Transcribe an audio file and store it in the DB."""
    filename = os.path.basename(audio_path)

    # Skip if already processed
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("SELECT transcript FROM transcripts WHERE filename=?", (filename,))
    row = c.fetchone()
    if row:
        conn.close()
        return row[0]

    # Transcribe
    result = model.transcribe(audio_path)
    transcript_text = result.get("text", "")

    # Store in DB
    c.execute(
        "INSERT INTO transcripts (filename, transcript) VALUES (?, ?)",
        (filename, transcript_text)
    )
    conn.commit()
    conn.close()
    return transcript_text

def export_to_docx(output_file="All_Transcripts.docx"):
    from docx import Document
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("SELECT filename, transcript FROM transcripts")
    rows = c.fetchall()
    conn.close()

    doc = Document()
    doc.add_heading("Call Transcriptions", level=0)
    for filename, transcript in rows:
        doc.add_heading(f"Audio File: {filename}", level=1)
        doc.add_paragraph(transcript if transcript else "N/A")
        doc.add_paragraph("\n")
    doc.save(output_file)
    return output_file
